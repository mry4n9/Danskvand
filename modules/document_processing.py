# modules/document_processing.py
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io

def create_transparency_document(
    company_url: str,
    website_text: str | None,
    website_summary: str | None,
    additional_text: str | None,
    additional_summary: str | None,
    lead_magnet_text: str | None,
    lead_magnet_summary: str | None
) -> bytes:
    """
    Creates a Word document containing extracted texts and their summaries.
    """
    doc = Document()

    def add_section(title, text_content, summary_content, text_header, summary_header):
        if text_content or summary_content:
            doc.add_heading(title, level=1)
            
            if text_content:
                doc.add_heading(text_header, level=2)
                p_text = doc.add_paragraph(str(text_content)[:90000]) # Limit length to prevent huge docs
                p_text.alignment = WD_ALIGN_PARAGRAPH.LEFT
                doc.add_paragraph() # Add some space
            
            if summary_content:
                doc.add_heading(summary_header, level=2)
                p_summary = doc.add_paragraph(str(summary_content))
                p_summary.alignment = WD_ALIGN_PARAGRAPH.LEFT
                doc.add_paragraph() # Add some space
            doc.add_page_break()

    doc.add_heading("AI Ad Copy Generation - Context Transparency Report", level=0)
    doc.add_paragraph(f"This document outlines the source texts and AI-generated summaries used for creating ad copy for the company associated with: {company_url}")
    doc.add_paragraph()

    add_section(
        "Client Website Context",
        website_text, website_summary,
        "Company URL Extract (Raw Text)", "Company URL Summarization (AI)"
    )
    
    add_section(
        "Additional Company Context",
        additional_text, additional_summary,
        "Additional Company Extract (Raw Text)", "Additional Company Summarization (AI)"
    )

    add_section(
        "Lead Magnet Context",
        lead_magnet_text, lead_magnet_summary,
        "Lead Magnet Extract (Raw Text)", "Lead Magnet Summarization (AI)"
    )

    # Save to a BytesIO object
    doc_bytes = io.BytesIO()
    doc.save(doc_bytes)
    doc_bytes.seek(0)
    return doc_bytes.getvalue()